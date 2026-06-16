"""Convert docs/MANUAL.md to a Google Docs-friendly .docx.

Adds a title/front page and an auto-TOC placeholder page, and applies real
Heading styles + tables so Google Docs can build a clickable Table of Contents
(Insert > Table of contents) after upload.
"""
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(ROOT, "docs", "MANUAL.md")
OUT = os.path.join(ROOT, "docs", "MANUAL.docx")

TEAL = RGBColor(0x13, 0x4E, 0x4A)


def add_inline(paragraph, text):
    """Render **bold**, `code`, and [text](url) within a paragraph."""
    # Strip link syntax to just the visible text.
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # Tokenize on bold and inline-code markers.
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def build_title_page(doc):
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("LGU Queuing System")
    r.bold = True
    r.font.size = Pt(34)
    r.font.color.rgb = TEAL

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("User Manual")
    rs.font.size = Pt(20)
    rs.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rsub = sub.add_run("Queue Management System for Local Government Units")
    rsub.italic = True
    rsub.font.size = Pt(12)

    for _ in range(10):
        doc.add_paragraph()

    for label in ("Prepared by: ______________________________",
                  "Office / LGU: ______________________________",
                  "Date: ______________________________"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(label).font.size = Pt(11)

    doc.add_page_break()


def build_toc_page(doc):
    h = doc.add_paragraph()
    r = h.add_run("Table of Contents")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = TEAL
    note = doc.add_paragraph()
    rn = note.add_run(
        "(In Google Docs: place the cursor below, then Insert > Table of "
        "contents. It builds automatically from the headings in this document.)"
    )
    rn.italic = True
    rn.font.size = Pt(9)
    rn.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()
    doc.add_page_break()


def parse_table(lines, i):
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(lines[i].strip())
        i += 1
    # rows[1] is the |---|---| separator
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    header = cells[0]
    body = cells[2:] if len(cells) > 2 else []
    return header, body, i


def main():
    with open(SRC, encoding="utf-8") as f:
        md = f.read()
    lines = md.split("\n")

    doc = Document()
    # Base font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    build_title_page(doc)
    build_toc_page(doc)

    i = 0
    skip_toc_block = False
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]

        # Code fences
        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.3)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Drop the markdown "Table of Contents" list block (GDocs builds its own)
        if stripped == "## Table of Contents":
            skip_toc_block = True
            i += 1
            continue
        if skip_toc_block:
            if stripped == "---":
                skip_toc_block = False
            i += 1
            continue

        # Horizontal rule -> skip
        if stripped == "---":
            i += 1
            continue

        # Title (single #) -> already on front page, skip
        if stripped.startswith("# ") and not stripped.startswith("## "):
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(re.sub(r"^###\s+", "", stripped), level=2)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(re.sub(r"^##\s+", "", stripped), level=1)
            i += 1
            continue

        # Tables
        if stripped.startswith("|"):
            header, body, i = parse_table(lines, i)
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for c, text in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.paragraphs[0].text = ""
                run = cell.paragraphs[0].add_run(re.sub(r"[*`]", "", text))
                run.bold = True
            for row in body:
                cells = table.add_row().cells
                for c, text in enumerate(row[: len(header)]):
                    cells[c].paragraphs[0].text = ""
                    add_inline(cells[c].paragraphs[0], text)
            doc.add_paragraph()
            continue

        # Blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run_text = stripped.lstrip(">").strip()
            add_inline(p, run_text)
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\d+\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m.group(1))
            i += 1
            continue

        # Blank
        if stripped == "":
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(OUT)
    print("Saved:", OUT)


if __name__ == "__main__":
    main()
